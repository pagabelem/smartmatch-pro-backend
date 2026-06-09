"""
Tests Phase 7 — Module Resumes (Upload CV)
Couvre : upload PDF/DOCX, validation extension/taille, CRUD, texte brut, permissions, extraction

Deux modes de test :
- Tests asynchrones (httpx AsyncClient) pour les routes API
- Tests unitaires synchrones pour les fonctions internes
"""

import io
import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock
from httpx import AsyncClient
from fastapi.testclient import TestClient

# Suppression du 'pytestmark = pytest.mark.asyncio' global pour éviter 
# de marquer à tort les classes de tests purement synchrones.


# ===========================================================================
# Helpers
# ===========================================================================

def make_pdf(content: bytes = b"%PDF-1.4 fake pdf content for testing") -> tuple:
    return ("cv.pdf", io.BytesIO(content), "application/pdf")


def make_docx(content: bytes = b"PK fake docx content") -> tuple:
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return ("cv.docx", io.BytesIO(content), mime)


def make_file(filename: str, content: bytes, mime: str) -> tuple:
    return (filename, io.BytesIO(content), mime)


def make_upload_file(filename: str, content: bytes, content_type: str):
    """Helper pour créer un fichier multipart pour httpx."""
    return ("file", (filename, io.BytesIO(content), content_type))


# ===========================================================================
# Tests d'upload asynchrones
# ===========================================================================

@pytest.mark.asyncio
class TestUploadResumeAsync:
    """Tests d'upload de CV asynchrones."""

    async def test_upload_pdf_success(
        self, async_client: AsyncClient, test_user, test_profile
    ):
        with patch("app.modules.resumes.resume_service.extract_raw_text", return_value="Extracted text"):
            resp = await async_client.post(
                "/api/v1/resumes/upload",
                files={"file": make_pdf()},
                headers=test_user["headers"],
            )
        assert resp.status_code in (200, 201)
        data = resp.json()
        resume_data = data.get("data", data)
        assert resume_data["filename"].endswith(".pdf")
        assert "id" in resume_data
        assert resume_data.get("is_parsed") in (True, False)

    async def test_upload_docx_success(
        self, async_client: AsyncClient, test_user, test_profile
    ):
        with patch("app.modules.resumes.resume_service.extract_raw_text", return_value="Resume text"):
            resp = await async_client.post(
                "/api/v1/resumes/upload",
                files={"file": make_docx()},
                headers=test_user["headers"],
            )
        assert resp.status_code in (200, 201)
        data = resp.json()
        resume_data = data.get("data", data)
        assert resume_data["filename"].endswith(".docx")

    async def test_upload_invalid_extension(
        self, async_client: AsyncClient, test_user, test_profile
    ):
        resp = await async_client.post(
            "/api/v1/resumes/upload",
            files={"file": make_file("cv.exe", b"malicious", "application/octet-stream")},
            headers=test_user["headers"],
        )
        assert resp.status_code in (400, 415, 422)

    async def test_upload_too_large(
        self, async_client: AsyncClient, test_user, test_profile
    ):
        # Génère un fichier de 6 Mo (dépasse la limite de 5 Mo)
        big_content = b"%PDF " + b"x" * (6 * 1024 * 1024)
        resp = await async_client.post(
            "/api/v1/resumes/upload",
            files={"file": make_file("big.pdf", big_content, "application/pdf")},
            headers=test_user["headers"],
        )
        assert resp.status_code in (400, 413, 422)

    async def test_upload_unauthenticated(self, async_client: AsyncClient):
        resp = await async_client.post(
            "/api/v1/resumes/upload",
            files={"file": make_pdf()},
        )
        assert resp.status_code == 401

    async def test_upload_no_file(
        self, async_client: AsyncClient, test_user, test_profile
    ):
        resp = await async_client.post(
            "/api/v1/resumes/upload",
            headers=test_user["headers"],
        )
        assert resp.status_code == 422

    async def test_upload_stores_relative_path(
        self, async_client: AsyncClient, test_user, test_profile
    ):
        """Le file_path stocké ne doit pas être un chemin absolu."""
        with patch("app.modules.resumes.resume_service.extract_raw_text", return_value="text"):
            resp = await async_client.post(
                "/api/v1/resumes/upload",
                files={"file": make_pdf()},
                headers=test_user["headers"],
            )
        assert resp.status_code in (200, 201)
        data = resp.json()
        resume_data = data.get("data", data)
        file_path = resume_data.get("file_path", "")
        assert not file_path.startswith("/"), "Le chemin ne doit pas être absolu"

    async def test_upload_multiple_resumes_same_profile(
        self, async_client: AsyncClient, test_user, test_profile
    ):
        with patch("app.modules.resumes.resume_service.extract_raw_text", return_value="text"):
            for i in range(3):
                resp = await async_client.post(
                    "/api/v1/resumes/upload",
                    files={"file": make_pdf(f"PDF {i}".encode())},
                    headers=test_user["headers"],
                )
                assert resp.status_code in (200, 201)


# ===========================================================================
# Tests de lecture d'un CV asynchrones
# ===========================================================================

@pytest.mark.asyncio
class TestGetResumeAsync:
    """Tests de récupération d'un CV par ID."""

    async def test_get_resume_by_id(
        self, async_client: AsyncClient, test_user, test_resume
    ):
        resume_id = test_resume["id"]
        resp = await async_client.get(
            f"/api/v1/resumes/{resume_id}",
            headers=test_user["headers"],
        )
        assert resp.status_code == 200
        data = resp.json()
        resume_data = data.get("data", data)
        assert resume_data["id"] == resume_id

    async def test_get_resume_not_found(self, async_client: AsyncClient, test_user):
        resp = await async_client.get(
            "/api/v1/resumes/99999999",
            headers=test_user["headers"],
        )
        assert resp.status_code == 404

    async def test_get_resume_other_user_forbidden(
        self, async_client: AsyncClient, test_second_user, test_resume
    ):
        resume_id = test_resume["id"]
        resp = await async_client.get(
            f"/api/v1/resumes/{resume_id}",
            headers=test_second_user["headers"],
        )
        assert resp.status_code == 403

    async def test_get_resume_admin_access(
        self, async_client: AsyncClient, test_admin, test_resume
    ):
        resume_id = test_resume["id"]
        resp = await async_client.get(
            f"/api/v1/resumes/{resume_id}",
            headers=test_admin["headers"],
        )
        assert resp.status_code == 200

    async def test_get_resume_unauthenticated(
        self, async_client: AsyncClient, test_resume
    ):
        resume_id = test_resume["id"]
        resp = await async_client.get(f"/api/v1/resumes/{resume_id}")
        assert resp.status_code == 401


# ===========================================================================
# Tests de récupération du texte brut
# ===========================================================================

@pytest.mark.asyncio
class TestGetResumeTextAsync:
    """Tests de récupération du texte brut d'un CV."""

    async def test_get_resume_text(
        self, async_client: AsyncClient, test_user, test_resume
    ):
        resume_id = test_resume["id"]
        resp = await async_client.get(
            f"/api/v1/resumes/{resume_id}/text",
            headers=test_user["headers"],
        )
        assert resp.status_code == 200
        data = resp.json()
        resume_data = data.get("data", data)
        assert "raw_text" in resume_data

    async def test_get_resume_text_other_user_forbidden(
        self, async_client: AsyncClient, test_second_user, test_resume
    ):
        resume_id = test_resume["id"]
        resp = await async_client.get(
            f"/api/v1/resumes/{resume_id}/text",
            headers=test_second_user["headers"],
        )
        assert resp.status_code == 403


# ===========================================================================
# Tests de liste des CVs d'un profil
# ===========================================================================

@pytest.mark.asyncio
class TestListResumesByProfileAsync:
    """Tests de liste paginée des CVs d'un profil."""

    async def test_get_resumes_by_profile(
        self, async_client: AsyncClient, test_user, test_profile, test_resume
    ):
        profile_id = test_profile["id"]
        resp = await async_client.get(
            f"/api/v1/resumes/profile/{profile_id}",
            headers=test_user["headers"],
        )
        assert resp.status_code == 200
        data = resp.json()
        list_data = data.get("data", data)
        items = list_data.get("items", list_data) if isinstance(list_data, dict) else list_data
        assert len(items) >= 1

    async def test_get_resumes_by_profile_paginated(
        self, async_client: AsyncClient, test_user, test_profile, test_resume
    ):
        profile_id = test_profile["id"]
        resp = await async_client.get(
            f"/api/v1/resumes/profile/{profile_id}?page=1&limit=5",
            headers=test_user["headers"],
        )
        assert resp.status_code == 200

    async def test_get_resumes_other_profile_forbidden(
        self, async_client: AsyncClient, test_second_user, test_profile
    ):
        profile_id = test_profile["id"]
        resp = await async_client.get(
            f"/api/v1/resumes/profile/{profile_id}",
            headers=test_second_user["headers"],
        )
        assert resp.status_code == 403

    async def test_get_resumes_profile_not_found(
        self, async_client: AsyncClient, test_admin
    ):
        resp = await async_client.get(
            "/api/v1/resumes/profile/99999999",
            headers=test_admin["headers"],
        )
        assert resp.status_code == 404


# ===========================================================================
# Tests de suppression d'un CV
# ===========================================================================

@pytest.mark.asyncio
class TestDeleteResumeAsync:
    """Tests de suppression d'un CV."""

    async def test_delete_resume_owner(
        self, async_client: AsyncClient, test_user, test_resume
    ):
        resume_id = test_resume["id"]
        with patch("pathlib.Path.exists", return_value=True):
            resp = await async_client.delete(
                f"/api/v1/resumes/{resume_id}",
                headers=test_user["headers"],
            )
        assert resp.status_code == 200

        # Vérification : le CV n'est plus accessible
        get_resp = await async_client.get(
            f"/api/v1/resumes/{resume_id}",
            headers=test_user["headers"],
        )
        assert get_resp.status_code == 404

    async def test_delete_resume_forbidden(
        self, async_client: AsyncClient, test_second_user, test_resume
    ):
        resume_id = test_resume["id"]
        resp = await async_client.delete(
            f"/api/v1/resumes/{resume_id}",
            headers=test_second_user["headers"],
        )
        assert resp.status_code == 403

    async def test_delete_resume_admin(
        self, async_client: AsyncClient, test_admin, test_resume
    ):
        resume_id = test_resume["id"]
        with patch("pathlib.Path.exists", return_value=True):
            resp = await async_client.delete(
                f"/api/v1/resumes/{resume_id}",
                headers=test_admin["headers"],
            )
        assert resp.status_code == 200

    async def test_delete_resume_not_found(
        self, async_client: AsyncClient, test_user
    ):
        resp = await async_client.delete(
            "/api/v1/resumes/99999999",
            headers=test_user["headers"],
        )
        assert resp.status_code == 404

    async def test_delete_resume_unauthenticated(
        self, async_client: AsyncClient, test_resume
    ):
        resume_id = test_resume["id"]
        resp = await async_client.delete(f"/api/v1/resumes/{resume_id}")
        assert resp.status_code == 401


# ===========================================================================
# Tests unitaires d'extraction de texte (PUREMENT SYNCHRONES)
# ===========================================================================

class TestExtractTextUnit:
    """Tests unitaires pour l'extraction de texte (synchrones)."""

    def test_extract_text_from_pdf(self, tmp_path):
        from app.modules.resumes.file_parser import parse_resume_file

        try:
            import pypdf
            from pypdf import PdfWriter

            pdf_path = tmp_path / "test.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=200, height=200)
            with open(pdf_path, "wb") as f:
                writer.write(f)

            result = parse_resume_file(str(pdf_path), "application/pdf")
            assert isinstance(result, str)
        except ImportError:
            pytest.skip("pypdf non installé")

    def test_extract_text_from_docx(self, tmp_path):
        from app.modules.resumes.file_parser import parse_resume_file

        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Python Developer")
            doc.add_paragraph("FastAPI · SQLAlchemy · React")
            docx_path = tmp_path / "test.docx"
            doc.save(str(docx_path))

            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            result = parse_resume_file(str(docx_path), mime)

            assert "Python Developer" in result
            assert "FastAPI" in result
        except ImportError:
            pytest.skip("python-docx non installé")

    def test_extract_text_unsupported_format(self, tmp_path):
        from app.modules.resumes.file_parser import parse_resume_file

        txt_path = tmp_path / "cv.txt"
        txt_path.write_text("hello")

        with pytest.raises(ValueError, match="non supporté"):
            parse_resume_file(str(txt_path), "text/plain")

    def test_extract_text_file_not_found(self):
        from app.modules.resumes.file_parser import parse_resume_file

        with pytest.raises(FileNotFoundError):
            parse_resume_file("/nonexistent/path/cv.pdf", "application/pdf")


# ===========================================================================
# Tests unitaires des services (PUREMENT SYNCHRONES)
# ===========================================================================

class TestResumeServiceUnit:
    """Tests unitaires pour les fonctions de service."""

    def test_service_validate_extension(self):
        from app.modules.resumes.resume_service import _validate_file
        from app.core.exceptions import ValidationException

        # ✅ Pattern (?i) pour ignorer la casse car le message lève 'Extension' (E majuscule)
        with pytest.raises(ValidationException, match="(?i)extension"):
            _validate_file("cv.exe", 100)

    def test_service_validate_size(self):
        from app.modules.resumes.resume_service import _validate_file
        from app.core.exceptions import ValidationException
        from app.config import settings

        oversized = settings.max_file_size_bytes + 1
        with pytest.raises(ValidationException, match="volumineux"):
            _validate_file("cv.pdf", oversized)


# ===========================================================================
# Tests avec client synchrone (TestClient FastAPI) pour compatibilité
# ===========================================================================

class TestResumeSync:
    """Tests synchrones avec TestClient FastAPI."""

    def test_upload_pdf_success_sync(self, sync_client: TestClient, auth_headers):
        with patch("app.modules.resumes.resume_service.extract_raw_text", return_value="Extracted text"):
            response = sync_client.post(
                "/api/v1/resumes/upload",
                files={"file": ("cv_test.pdf", io.BytesIO(b"%PDF-1.4 fake"), "application/pdf")},
                headers=auth_headers,
            )
        assert response.status_code in (200, 201)

    def test_upload_unauthenticated_sync(self, sync_client: TestClient):
        response = sync_client.post(
            "/api/v1/resumes/upload",
            files={"file": ("cv.pdf", io.BytesIO(b"%PDF fake"), "application/pdf")},
        )
        assert response.status_code == 401